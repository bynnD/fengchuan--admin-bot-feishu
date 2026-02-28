"""
通用文件内容提取模块 - 供所有有附件识别读取需求的工单类型使用。
支持 Word/PDF/Excel 文本提取，图片和扫描件通过 RapidOCR 本地识别。
"""

import logging
import numpy as np

logger = logging.getLogger(__name__)

# RapidOCR 单例，首次调用时懒加载
_rapid_ocr = None


def _get_rapid_ocr():
    """懒加载 RapidOCR，避免启动时加载模型。"""
    global _rapid_ocr
    if _rapid_ocr is None:
        from rapidocr_onnxruntime import RapidOCR
        _rapid_ocr = RapidOCR()
    return _rapid_ocr


def rapid_ocr(image_content):
    """使用 RapidOCR 识别图片/扫描件中的文字。返回拼接后的文本。"""
    if not image_content:
        return ""
    try:
        import cv2
        from PIL import Image
        from io import BytesIO
        img = Image.open(BytesIO(image_content))
        # RGBA/透明图合成白底，避免透明区域变黑影响识别
        if img.mode in ("RGBA", "P"):
            background = Image.new("RGB", img.size, (255, 255, 255))
            if img.mode == "RGBA":
                background.paste(img, mask=img.split()[3])
            else:
                background.paste(img)
            img = background
        elif img.mode != "RGB":
            img = img.convert("RGB")
        img_array = np.array(img)
        img_array = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
        # 小图放大，提升截图/手机拍图的识别率
        h, w = img_array.shape[:2]
        if max(h, w) < 1600:
            scale = 1600 / max(h, w)
            new_size = (int(w * scale), int(h * scale))
            img_array = cv2.resize(img_array, new_size, interpolation=cv2.INTER_LANCZOS4)
        ocr = _get_rapid_ocr()
        result, _ = ocr(img_array)
        if not result:
            return ""
        texts = [item[1] for item in result if item and len(item) >= 2 and item[1]]
        return "\n".join(t for t in texts if t) if texts else ""
    except Exception as e:
        logger.warning("RapidOCR 异常: %s", e)
        return ""


def extract_text_from_file(file_content, file_name, get_token):
    """
    从 Word/PDF/图片/扫描件提取文本。
    图片和扫描件使用 RapidOCR 本地识别。返回前 8000 字符。
    供所有有附件识别读取需求的工单类型调用。
    get_token 保留以兼容调用方，本地 OCR 不使用。
    """
    if not file_content:
        return ""
    ext = (file_name.rsplit(".", 1)[-1] or "").lower()
    try:
        if ext in ("docx", "doc"):
            from docx import Document
            from io import BytesIO
            try:
                doc = Document(BytesIO(file_content))
                parts = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                parts.append(cell.text)
                text = "\n".join(parts)[:8000]
                if not text.strip():
                    logger.debug("Word 解析(%s): 段落和表格均为空", file_name)
                return text
            except Exception as doc_err:
                logger.warning("Word 解析失败(%s): %s", file_name, doc_err)
                return ""
        if ext == "xlsx":
            from io import BytesIO
            try:
                from openpyxl import load_workbook
                wb = load_workbook(BytesIO(file_content), read_only=True, data_only=True)
                parts = []
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        row_cells = [str(c).strip() for c in (row or []) if c is not None and str(c).strip()]
                        if row_cells:
                            parts.append(" | ".join(row_cells))
                text = "\n".join(parts)[:8000]
                return text
            except Exception as excel_err:
                logger.warning("Excel 解析失败(%s): %s", file_name, excel_err)
                return ""
        if ext in ("png", "jpg", "jpeg", "bmp", "gif", "webp"):
            return rapid_ocr(file_content)[:8000]
        if ext == "pdf":
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(file_content))
            parts = []
            for page in reader.pages[:20]:
                t = page.extract_text()
                if t:
                    parts.append(t)
            text = "\n".join(parts)
            if len(text.strip()) >= 50:
                return text[:8000]
            # 文本很少，可能是扫描件，用 PyMuPDF 渲染页面为图片后 OCR
            import fitz
            doc = fitz.open(stream=file_content, filetype="pdf")
            ocr_parts = []
            for i in range(min(len(doc), 20)):
                page = doc[i]
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                ocr_text = rapid_ocr(img_bytes)
                if ocr_text:
                    ocr_parts.append(ocr_text)
            doc.close()
            return "\n".join(ocr_parts)[:8000] if ocr_parts else text[:8000]
    except Exception as e:
        logger.warning("提取文件文本失败(%s): %s", file_name, e)
    return ""
